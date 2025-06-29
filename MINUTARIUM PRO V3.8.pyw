import os
import json
import shutil
import tempfile
import re
import webbrowser
import customtkinter as ctk
from tkinter import messagebox, Listbox
from tkinterdnd2 import TkinterDnD, DND_FILES
from docx import Document
from docx.shared import Pt
import atexit

try:
    import winreg
except ImportError:
    winreg = None

try:
    import win32com.client as win32
except ImportError:
    messagebox.showerror("Erro de Dependência", "O módulo 'pywin32' não está instalado. Execute 'pip install pywin32' para instalá-lo.")
    raise

# --- CONFIGURAÇÕES ---

DB_FILE_PATH = r"G:\Meu Drive\CONTROLLER\DATA CENTER\BANCO_DE_DADOS.txt"

NEON_TRON_THEME = {
    "CTk": {"fg_color": ["#EAEAEA", "#000000"]},
    "CTkToplevel": {"fg_color": ["#EAEAEA", "#000000"]},
    "CTkFrame": {"corner_radius": 6, "border_width": 0, "fg_color": ["#DCDCDC", "#0A0A0A"], "top_fg_color": ["#C8C8C8", "#050505"], "border_color": ["#B0B0B0", "#007080"]},
    "CTkButton": {"corner_radius": 6, "border_width": 0, "fg_color": ["#00D1D1", "#00D1D1"], "hover_color": ["#00A0A0", "#00A0A0"], "border_color": ["#00D1D1", "#00D1D1"], "text_color": ["#001010", "#001010"], "text_color_disabled": ["#007080", "#007080"]},
    "CTkLabel": {"corner_radius": 0, "fg_color": "transparent", "text_color": ["#424242", "#00D1D1"]},
    "CTkEntry": {"corner_radius": 6, "border_width": 1, "fg_color": ["#FFFFFF", "#1A1A1A"], "border_color": ["#B0B0B0", "#00D1D1"], "text_color": ["#333333", "#E0E0E0"], "placeholder_text_color": ["#9E9E9E", "#00A0A0"]},
    "CTkTextbox": {"corner_radius": 6, "border_width": 1, "fg_color": ["#FFFFFF", "#0A0A0A"], "border_color": ["#B0B0B0", "#00D1D1"], "text_color": ["#333333", "#E0E0E0"], "scrollbar_button_color": ["#BDBDBD", "#505050"], "scrollbar_button_hover_color": ["#ADADAD", "#00A0A0"]},
    "CTkFont": {"family": "Helvetica", "size": 12, "weight": "normal"},
    "CTkComboBox": {"corner_radius": 6, "border_width": 2, "fg_color": ["#F9F9FA", "#343638"], "border_color": ["#979DA2", "#565B5E"], "button_color": ["#979DA2", "#565B5E"], "button_hover_color": ["#6E7174", "#7A848D"], "text_color": ["#000000", "#FFFFFF"], "text_color_disabled": ["#AAB0B5", "#7A848D"]},
    "CTkSlider": {"corner_radius": 1000, "button_corner_radius": 1000, "border_width": 4, "button_length": 0, "fg_color": ["#979DA2", "#565B5E"], "progress_color": ["#6E7174", "#7A848D"], "button_color": ["#3B8ED0", "#1F6AA5"], "button_hover_color": ["#36719F", "#144870"]},
    "CTkSegmentedButton": {"corner_radius": 6, "border_width": 2, "fg_color": ["#979DA2", "#565B5E"], "selected_color": ["#3B8ED0", "#1F6AA5"], "selected_hover_color": ["#36719F", "#144870"], "unselected_color": ["#979DA2", "#565B5E"], "unselected_hover_color": ["#6E7174", "#7A848D"], "text_color": ["#FFFFFF", "#FFFFFF"], "text_color_disabled": ["#AAB0B5", "#7A848D"]},
    "CTkProgressBar": {"corner_radius": 1000, "border_width": 0, "fg_color": ["#979DA2", "#565B5E"], "progress_color": ["#3B8ED0", "#1F6AA5"], "border_color": ["#6E7174", "#565B5E"]},
    "CTkScrollbar": {"corner_radius": 1000, "border_spacing": 4, "fg_color": "transparent", "button_color": ["#979DA2", "#565B5E"], "button_hover_color": ["#6E7174", "#7A848D"]},
    "CTkOptionMenu": {"corner_radius": 6, "fg_color": ["#3B8ED0", "#1F6AA5"], "button_color": ["#36719F", "#144870"], "button_hover_color": ["#27577D", "#103957"], "text_color": ["#FFFFFF", "#FFFFFF"], "text_color_disabled": ["#AAB0B5", "#7A848D"]},
    "CTkCheckBox": {"corner_radius": 6, "border_width": 3, "fg_color": ["#3B8ED0", "#1F6AA5"], "border_color": ["#3B8ED0", "#1F6AA5"], "hover_color": ["#36719F", "#144870"], "checkmark_color": ["#FFFFFF", "#FFFFFF"], "text_color": ["#000000", "#FFFFFF"], "text_color_disabled": ["#AAB0B5", "#7A848D"]},
    "CTkSwitch": {"corner_radius": 1000, "border_width": 3, "button_length": 0, "fg_color": ["#979DA2", "#565B5E"], "progress_color": ["#3B8ED0", "#1F6AA5"], "button_color": ["#F9F9FA", "#2B2B2B"], "button_hover_color": ["#C2C2C2", "#424242"], "text_color": ["#000000", "#FFFFFF"], "text_color_disabled": ["#AAB0B5", "#7A848D"]},
    "CTkRadioButton": {"corner_radius": 1000, "border_width_checked": 6, "border_width_unchecked": 3, "fg_color": ["#3B8ED0", "#1F6AA5"], "border_color": ["#3B8ED0", "#1F6AA5"], "hover_color": ["#36719F", "#144870"], "text_color": ["#000000", "#FFFFFF"], "text_color_disabled": ["#AAB0B5", "#7A848D"]},
    "DropdownMenu": {"fg_color": ["#F9F9FA", "#2B2B2B"], "hover_color": ["#3B8ED0", "#1F6AA5"], "text_color": ["#000000", "#FFFFFF"]}
}

def carregar_dicionario_externo(caminho_arquivo):
    db_directory = os.path.dirname(caminho_arquivo)
    os.makedirs(db_directory, exist_ok=True)

    if not os.path.exists(caminho_arquivo):
        messagebox.showerror("Erro Crítico", f"O arquivo de banco de dados não foi encontrado no caminho:\n{caminho_arquivo}\n\nO diretório foi criado, mas o arquivo precisa ser colocado nele.")
        return None
    
    try:
        with open(caminho_arquivo, 'r', encoding='utf-8') as file:
            conteudo_bruto = file.read()
        
        conteudo_sanitizado = re.sub(r',\s*\}', '}', conteudo_bruto)
        
        dicionario = json.loads(conteudo_sanitizado)
        return dicionario
        
    except json.JSONDecodeError as e:
        mensagem_erro = (
            f"Falha ao ler ou processar o arquivo de banco de dados:\n\n"
            f"Erro de formato JSON: {e}\n\n"
            f"CAUSA PROVÁVEL: Existe um erro de sintaxe no arquivo '{os.path.basename(caminho_arquivo)}'.\n"
            f"Verifique se todas as chaves e valores estão entre aspas duplas \" e se as linhas terminam com vírgula (,), exceto a última."
        )
        messagebox.showerror("Erro Crítico de Formato", mensagem_erro)
        return None
    except Exception as e:
        messagebox.showerror("Erro Crítico", f"Ocorreu um erro inesperado ao processar o banco de dados:\n{e}")
        return None

# --- LÓGICA DE NEGÓCIO (Inalterada) ---

def convert_files(files):
    script_directory = os.path.dirname(os.path.abspath(__file__))
    converted_folder = os.path.join(script_directory, "BIPAGEM CONVERTIDA")
    os.makedirs(converted_folder, exist_ok=True)

    converted_files = []
    for file_path in files:
        converted_file_path = os.path.join(converted_folder, os.path.basename(file_path))
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                lines = file.readlines()
        except Exception as e:
            messagebox.showerror("Erro", f"Erro ao ler o arquivo {file_path}:\n{e}")
            continue

        try:
            with open(converted_file_path, 'w', encoding='utf-8') as converted_file:
                for line in lines:
                    parts = line.strip().split()
                    if len(parts) == 2:
                        code, quantity = parts
                        item = substituicoes_ean_para_nx.get(code, code)
                        converted_file.write(f"{item} {quantity}\n")
                    else:
                        converted_file.write(line)
        except Exception as e:
            messagebox.showerror("Erro", f"Erro ao escrever o arquivo convertido {converted_file_path}:\n{e}")
            continue
        converted_files.append(converted_file_path)
    return converted_files

def natural_sort_key(s):
    filename = os.path.basename(s)
    return [int(text) if text.isdigit() else text.lower() for text in re.split(r'(\d+)', filename)]

def process_files(destination, order_number, files):
    base_path = r'G:\Meu Drive\CONTROLLER\MINUTA'
    os.makedirs(base_path, exist_ok=True)

    client_folder = None
    for folder_name in os.listdir(base_path):
        if destination.lower() == folder_name.lower():
            client_folder = os.path.join(base_path, folder_name)
            break
    
    if not client_folder:
        client_folder = os.path.join(base_path, destination)
        os.makedirs(client_folder, exist_ok=True)

    template_path = r"G:\Meu Drive\CONTROLLER\MINUTA\MODELO_MINUTA_TRANSPORTE_AUTOMATICA.docx"
    if not os.path.exists(template_path):
        messagebox.showerror("Erro", f"Template não encontrado: {template_path}")
        return None

    document = Document(template_path)
    for paragraph in document.paragraphs:
        if 'DESTINO:' in paragraph.text:
            paragraph.text = paragraph.text.replace('DESTINO:', f'DESTINO: {destination}')
            paragraph.runs[0].font.bold = True
        if 'PEDIDO:' in paragraph.text:
            paragraph.text = paragraph.text.replace('PEDIDO:', f'PEDIDO: {order_number}')
            paragraph.runs[0].font.bold = True

    for file_path in sorted(files, key=natural_sort_key):
        file_name = os.path.basename(file_path)
        file_name_no_ext = os.path.splitext(file_name)[0]
        lower_file_name = file_name_no_ext.lower()
        
        if 'saco' in lower_file_name or 'volume' in lower_file_name:
            words = file_name_no_ext.split()
            number = words[-1] if words[-1].isdigit() else ''
            section_title = f"SACO {number}" if number else "SACO"
        elif 'palete' in lower_file_name or 'pallet' in lower_file_name:
            words = file_name_no_ext.split()
            number = words[-1] if words[-1].isdigit() else ''
            section_title = f"PALETE {number}" if number else "PALETE"
        else:
            section_title = "SACO"

        title_paragraph = document.add_paragraph()
        title_run = title_paragraph.add_run(section_title)
        title_run.font.bold = True
        title_run.font.size = Pt(14)

        table = document.add_table(rows=2, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'DESCRIÇÃO'
        hdr_cells[1].text = 'QUANTIDADE'
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                run = paragraph.runs[0]
                run.font.bold = True
                paragraph.alignment = 0

        description_text = ""
        quantity_text = ""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                for line in file.readlines():
                    parts = line.strip().split()
                    if len(parts) >= 2:
                        quantity = parts[-1]
                        item = ' '.join(parts[:-1])
                        description_text += f"{item}\n"
                        quantity_text += f"{quantity}\n"
        except Exception as e:
            messagebox.showerror("Erro", f"Erro ao ler o arquivo convertido {file_path}:\n{e}")
            continue

        table.rows[1].cells[0].text = description_text.strip()
        table.rows[1].cells[1].text = quantity_text.strip()
        for cell in table.rows[1].cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = 0
                for run in paragraph.runs:
                    run.font.size = Pt(12)
        document.add_paragraph()

    docx_filename = f"MINUTA {destination} {order_number}.docx"
    pdf_filename = f"MINUTA {destination} {order_number}.pdf"
    docx_path = os.path.join(client_folder, docx_filename)
    pdf_path = os.path.join(client_folder, pdf_filename)

    try:
        document.save(docx_path)
    except Exception as e:
        messagebox.showerror("Erro", f"Erro ao salvar o arquivo Word:\n{e}")
        return None

    try:
        gen_py_path = os.path.join(tempfile.gettempdir(), 'gen_py')
        if os.path.isdir(gen_py_path):
            shutil.rmtree(gen_py_path)
    except Exception as e:
        print(f"Aviso: Não foi possível limpar o cache do pywin32. Erro: {e}")

    try:
        word = win32.gencache.EnsureDispatch('Word.Application')
        word.Visible = False
        word_doc = word.Documents.Open(docx_path)
        word_doc.SaveAs(pdf_path, FileFormat=17)
        word_doc.Close()
        word.Quit()
    except Exception as e:
        messagebox.showerror("Erro", f"Erro ao converter para PDF:\n{e}")
        return None

    try:
        os.remove(docx_path)
    except Exception as e:
        messagebox.showerror("Erro", f"Erro ao remover o arquivo Word temporário:\n{e}")

    messagebox.showinfo("SUCESSO!", f"Arquivo PDF salvo em:\n{pdf_path}")
    return pdf_path

def get_windows_theme():
    if winreg is None:
        return "light"
    try:
        key_path = r'Software\Microsoft\Windows\CurrentVersion\Themes\Personalize'
        key = winreg.OpenKey(winreg.HKEY_CURRENT_USER, key_path)
        value, _ = winreg.QueryValueEx(key, 'AppsUseLightTheme')
        winreg.CloseKey(key)
        return "light" if value > 0 else "dark"
    except FileNotFoundError:
        return "light"

# --- INTERFACE GRÁFICA (Versão Final) ---

class App(ctk.CTk, TkinterDnD.DnDWrapper):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.TkdndVersion = TkinterDnD._require(self)
        
        self.title("Minutarium Pro (Definitive Edition)")
        
        self.geometry("480x480")
        self.resizable(False, False)

        self.dragged_files = []
        self.default_folder_path = r'G:\Meu Drive\CONTROLLER\MINUTA'
        self.last_generated_path = None
        self.changelog_window = None

        # --- Widgets ---
        self.main_frame = ctk.CTkFrame(self)
        self.main_frame.grid(row=0, column=0, padx=10, pady=10, sticky="nsew")
        self.grid_rowconfigure(0, weight=1)
        self.grid_columnconfigure(0, weight=1)

        self.main_frame.grid_columnconfigure(1, weight=1)
        self.main_frame.grid_rowconfigure(2, weight=1)

        ctk.CTkLabel(self.main_frame, text="Nome do Cliente:").grid(row=0, column=0, padx=(10, 5), pady=(10, 5), sticky="e")
        self.destination_entry = ctk.CTkEntry(self.main_frame, placeholder_text="Ex: Cliente")
        self.destination_entry.grid(row=0, column=1, padx=(0, 10), pady=(10, 5), sticky="ew")

        ctk.CTkLabel(self.main_frame, text="Nº Doc.Saída ou NF:").grid(row=1, column=0, padx=(10, 5), pady=5, sticky="e")
        self.order_entry = ctk.CTkEntry(self.main_frame, placeholder_text="9999 ou NF 9.999")
        self.order_entry.grid(row=1, column=1, padx=(0, 10), pady=5, sticky="ew")

        # --- INÍCIO DA ALTERAÇÃO ---
        self.listbox_container = ctk.CTkFrame(self.main_frame, fg_color=("white", "#000000"))
        self.listbox_container.grid(row=2, column=0, columnspan=2, padx=10, pady=10, sticky="nsew")

        self.file_listbox = Listbox(self.listbox_container, selectmode="multiple", width=50, height=10,
                                    highlightthickness=0, borderwidth=0)
        self.file_listbox.pack(side="left", fill="both", expand=True)

        self.placeholder_label = ctk.CTkLabel(self.listbox_container, text="Arraste os arquivos do Scanner para cá...",
                                              text_color=("#808080", "#A9A9A9"), fg_color="transparent")
        self.placeholder_label.place(relx=0.5, rely=0.5, anchor="center")
        
        self.info_frame = ctk.CTkFrame(self.main_frame, fg_color="transparent")
        self.info_frame.grid(row=3, column=0, columnspan=2, padx=10, pady=(5, 0), sticky="ew")
        self.info_frame.grid_columnconfigure(0, weight=1)
        
        self.info_label = ctk.CTkLabel(self.info_frame, text="Atualizações (Ctrl+H) | Alterar Tema (Ctrl+M)", font=("Helvetica", 10, "italic"))
        self.info_label.grid(row=0, column=0)

        self.button_frame = ctk.CTkFrame(self.main_frame, fg_color="transparent")
        self.button_frame.grid(row=4, column=0, columnspan=2, pady=10)
        # --- FIM DA ALTERAÇÃO ---
        
        self.generate_button = ctk.CTkButton(self.button_frame, text="Gerar Minuta", command=self.generate_document)
        self.generate_button.pack(side="left", padx=5, expand=True)

        self.go_to_folder_button = ctk.CTkButton(self.button_frame, text="Ir para Pasta", command=self.open_folder, state="normal")
        self.go_to_folder_button.pack(side="left", padx=5, expand=True)

        self.reset_button = ctk.CTkButton(self.button_frame, text="Redefinir", command=self.reset_fields, fg_color="#D32F2F", hover_color="#B71C1C")
        self.reset_button.pack(side="left", padx=5, expand=True)

        self.status_frame = ctk.CTkFrame(self, fg_color="transparent")
        self.status_frame.grid(row=1, column=0, sticky="ew", padx=10, pady=(0, 5))
        self.status_frame.grid_columnconfigure(0, weight=1)
        self.status_label = ctk.CTkLabel(self.status_frame, text="BANCO DE DADOS CARREGADO COM SUCESSO !!!", text_color=("#2E7D32", "#66BB6A"), font=ctk.CTkFont(size=12))
        self.status_label.grid(row=0, column=0)

        self.file_listbox.drop_target_register(DND_FILES)
        self.file_listbox.dnd_bind('<<Drop>>', self.drop_files)
        
        self.update_listbox_theme()
        
        self.bind("<Control-h>", self.show_changelog)
        self.bind("<Control-H>", self.show_changelog)
        self.bind("<Control-m>", self.toggle_manual_theme)
        self.bind("<Control-M>", self.toggle_manual_theme)

    def show_changelog(self, event=None):
        if self.changelog_window is None or not self.changelog_window.winfo_exists():
            self.changelog_window = ctk.CTkToplevel(self)
            self.changelog_window.title("Histórico de Alterações")
            self.changelog_window.geometry("600x500")
            self.changelog_window.resizable(False, False)
            
            self.changelog_window.grid_columnconfigure(0, weight=1)
            self.changelog_window.grid_rowconfigure(1, weight=1)

            title_label = ctk.CTkLabel(self.changelog_window, text="Minutarium Pro", font=ctk.CTkFont(size=20, weight="bold"))
            title_label.grid(row=0, column=0, padx=10, pady=(10, 5))

            textbox = ctk.CTkTextbox(self.changelog_window, wrap="word")
            textbox.grid(row=1, column=0, sticky="nsew", padx=10, pady=(5, 10))

            changelog_text = """
## Versão 3.8 (Atual)
- Adicionada criação automática de diretórios para o banco de dados e para as minutas, caso não existam.
- Centralizado o texto do placeholder na caixa de arrastar e soltar.

## Versão 3.7
- Adicionada limpeza automática da pasta temporária 'BIPAGEM CONVERTIDA' após cada geração de minuta (bem-sucedida ou não).

## Versão 3.6
- Adicionado atalho (Ctrl+M) para alternar manualmente entre os modos Light e Dark.
- Reordenados os botões para melhor UX/UI.

## Versão 3.5
- A lista de arquivos agora exibe apenas o nome do arquivo, não o caminho completo.

## Versão 3.4
- Corrigida a aplicação do tema na Listbox, removendo as bordas e aplicando as cores corretas para uma integração visual perfeita.

## Versão 3.3
- Corrigido erro 'KeyError' ao embutir um tema incompleto. O dicionário completo do tema agora está no código, garantindo a inicialização correta.

## Versão 3.2
- Corrigido erro de inicialização 'NameError' ao não importar o módulo 'atexit'.

## Versão 3.1
- Corrigido erro ao aplicar tema personalizado embutido. O tema agora é carregado através de um arquivo temporário, mantendo o programa autocontido.

## Versão 3.0
- Refatoração do sistema de temas para ser autocontido e mais robusto.
- Adicionado tema personalizado "Neon Tron" para o modo escuro.
- Adicionado rodapé com botões na janela de changelog.
- Botão "Código Fonte" agora abre o link do GitHub.
- Atalho do changelog agora funciona com Ctrl+h e Ctrl+H.
- Adicionado texto de exemplo (placeholder) nos campos de entrada.
- Movido o texto "Arraste os arquivos..." para dentro da caixa de listagem.
- Adicionado informativo sobre o atalho do changelog na tela principal.

## Versão 2.4
- Removido o seletor de tema manual.
- O programa agora detecta e aplica automaticamente o tema do Windows (Light/Dark).
- Adicionada janela de Histórico de Alterações (CTRL+H).
- Centralizada a mensagem de status no rodapé.
- Definido um tamanho fixo para a janela (não redimensionável).
- Adicionado rodapé com status de carregamento do banco de dados.

## Versão 2.3
- Refatoração do sistema de banco de dados para suportar formato JSON, corrigindo múltiplos erros de parsing.
- Melhorada a mensagem de erro ao falhar o carregamento do banco de dados.

## Versão 2.2
- Adicionada confirmação de segurança ao apagar arquivos de origem.
- Corrigida a função "Redefinir" para apagar a pasta temporária.
- Corrigida a ordenação dos arquivos para ser numérica (SACO 2 antes de SACO 10).

## Versão 2.1
- Botão "Ir para Pasta" agora abre o diretório padrão se nenhuma minuta foi gerada.
- Removido espaço vazio abaixo dos botões para um layout mais compacto.
- Botão "Ir para Pasta" adicionado, habilitado após gerar minuta.
- Botões "Gerar" e "Redefinir" centralizados.

## Versão 2.0
- Grande refatoração da interface para CustomTkinter.
- Corrigido erro crítico 'CLSIDToClassMap' do pywin32 limpando o cache.
- Dicionário de conversão movido para um arquivo externo.

## Versão 1.0
- Versão inicial com a lógica principal de geração de minutas em Tkinter.
"""
            textbox.insert("0.0", changelog_text)
            textbox.configure(state="disabled")
            
            footer_frame = ctk.CTkFrame(self.changelog_window, fg_color="transparent")
            footer_frame.grid(row=2, column=0, sticky="ew", padx=10, pady=(0, 10))
            footer_frame.grid_columnconfigure((0, 3), weight=1) 

            github_button = ctk.CTkButton(footer_frame, text="Suporte / Sobre", command=lambda: webbrowser.open_new_tab("https://github.com/SanctusLocalhost"))
            github_button.grid(row=0, column=1, padx=5)

            close_button = ctk.CTkButton(footer_frame, text="Fechar", command=self.changelog_window.destroy)
            close_button.grid(row=0, column=2, padx=5)
            
        else:
            self.changelog_window.focus()

    def update_listbox_theme(self):
        theme = ctk.get_appearance_mode()
        if theme == "Dark":
            self.file_listbox.config(bg="#000000", 
                                     fg="#E0E0E0", 
                                     selectbackground="#00D1D1", 
                                     selectforeground="#001010")
        else: # Light
            self.file_listbox.config(bg="white", 
                                     fg="black", 
                                     selectbackground="#3B8ED0", 
                                     selectforeground="white")

    def drop_files(self, event):
        self.placeholder_label.place_forget()
        
        full_paths = self.tk.splitlist(event.data)
        
        for full_path in full_paths:
            if full_path not in self.dragged_files:
                file_name = os.path.basename(full_path)
                self.file_listbox.insert("end", file_name)
                self.dragged_files.append(full_path)

    def toggle_manual_theme(self, event=None):
        current_mode = ctk.get_appearance_mode()
        new_mode = "light" if current_mode == "Dark" else "dark"
        
        ctk.set_appearance_mode(new_mode)
        self.update_listbox_theme()

    def _cleanup_converted_folder(self):
        script_directory = os.path.dirname(os.path.abspath(__file__))
        converted_folder = os.path.join(script_directory, "BIPAGEM CONVERTIDA")
        if os.path.exists(converted_folder):
            try:
                shutil.rmtree(converted_folder)
            except Exception as e:
                print(f"Aviso: Falha ao limpar a pasta temporária {converted_folder}: {e}")

    def generate_document(self):
        destination = self.destination_entry.get()
        order_number = self.order_entry.get()
        files = self.dragged_files

        if not destination or not order_number or not files:
            messagebox.showerror("Erro", "Por favor, preencha todos os campos e selecione os arquivos.")
            return

        try:
            converted_files = convert_files(files)
            if not converted_files:
                messagebox.showerror("Erro", "Nenhum arquivo foi convertido.")
                return
            
            pdf_path = process_files(destination, order_number, converted_files)
            
            if pdf_path:
                self.last_generated_path = pdf_path
        finally:
            self._cleanup_converted_folder()

    def open_folder(self):
        if self.last_generated_path and os.path.exists(self.last_generated_path):
            path_to_open = os.path.dirname(self.last_generated_path)
        else:
            path_to_open = self.default_folder_path

        if os.path.isdir(path_to_open):
            os.startfile(path_to_open)
        else:
            messagebox.showwarning("Aviso", f"O diretório não foi encontrado:\n{path_to_open}")

    def reset_fields(self):
        delete_source_files = False
        if self.dragged_files:
            delete_source_files = messagebox.askyesno(
                "Atenção - Exclusão de Arquivos",
                "Deseja Apagar os Arquivos de origem, que foram arrastados?\n\nUma vez feito, não pode ser desfeito."
            )
        
        if delete_source_files:
            for file_path in self.dragged_files:
                try:
                    if os.path.exists(file_path):
                        os.remove(file_path)
                except Exception as e:
                    messagebox.showerror("Erro de Exclusão", f"Não foi possível apagar o arquivo:\n{file_path}\n\nErro: {e}")

        self.destination_entry.delete(0, "end")
        self.order_entry.delete(0, "end")
        self.file_listbox.delete(0, "end")
        
        self.placeholder_label.place(relx=0.5, rely=0.5, anchor="center")
        
        self.last_generated_path = None
        
        self._cleanup_converted_folder()
        
        self.dragged_files = []

if __name__ == "__main__":
    substituicoes_ean_para_nx = carregar_dicionario_externo(DB_FILE_PATH)
    
    if substituicoes_ean_para_nx is None:
        exit()
    
    with tempfile.NamedTemporaryFile(mode='w', delete=False, suffix=".json", encoding='utf-8') as f:
        json.dump(NEON_TRON_THEME, f, indent=2)
        temp_theme_file_path = f.name
    
    ctk.set_default_color_theme(temp_theme_file_path)
    
    atexit.register(os.remove, temp_theme_file_path)
    
    system_theme = get_windows_theme()
    ctk.set_appearance_mode(system_theme)
    
    app = App()
    app.mainloop()